import base64
import io
import time

import docx
import openai
import os
import pandas as pd
import streamlit as st
from docx import Document
from PIL import Image
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError

def split_text(self, text, width):
    lines = text.split("\n")
    wrapped_lines = [word for line in lines for word in line.split(" ")]
    return wrapped_lines

@st.cache_resource()
def read_pdf(file):
    try:
        pdf_reader = PdfReader(file)
    except PdfReadError:
        st.error("Unsupported PDF format")
        return ""

    text = ""
    for page_num in range(min(len(pdf_reader.pages), 5)):
        text += pdf_reader.pages[page_num].extract_text()
    return text

def read_txt(file):
    return file.read()

def read_docx(file):
    try:
        doc = docx.Document(file)
    except docx.opc.exceptions.PackageNotFoundError:
        st.error("Unsupported Word document format")
        return ""

    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_image(file):
    try:
        image = Image.open(file)
        st.image(image, caption="Uploaded Image", use_column_width=True)
        return None
    except Exception as e:
        st.error(f"An error occurred while reading the image: {e}")
        return None

@st.cache_data()
def generate_answer(prompt, temperature, max_tokens, top_p):
    try:
        response = openai.Completion.create(
            engine="text-davinci-002",
            prompt=prompt,
            max_tokens=max_tokens,
            n=1,
            stop=None,
            temperature=temperature,
            top_p=top_p,
        )
        return response.choices[0].text.strip()
    except (openai.error.InvalidRequestError, openai.error.AuthenticationError, openai.error.APIConnectionError,
            openai.error.APIError, openai.error.RateLimitError) as e:
        st.error(f"An error occurred while generating the answer: {e}")
        return ""

def process_uploaded_file(uploaded_file):
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    if file_extension == ".pdf":
        return read_pdf(uploaded_file)
    elif file_extension == ".docx":
        return read_docx(uploaded_file)
    elif file_extension == ".txt":
        return read_txt(uploaded_file)
    elif file_extension in [".png", ".jpg", ".jpeg"]:
        return read_image(uploaded_file)
    else:
        st.error("Unsupported file format")
        return None

def display_history_and_favorites():
    st.sidebar.write("History:")
    for idx, item in enumerate(st.session_state.history):
        st.sidebar.markdown(f"**Question {idx + 1}:** {item['question']}")
        st.sidebar.markdown(f"**Answer {idx + 1}:** {item['answer']}")
        unique_key = f"delete_history_{idx}_{time.time()}"
        delete_button = st.sidebar.button(f"Delete history item {idx + 1}", key=unique_key)
        if delete_button:
            st.session_state.history.pop(idx)
            display_history_and_favorites()
            break

    st.sidebar.write("Favorites:")
    for idx, item in enumerate(st.session_state.favorites):
        st.sidebar.markdown(f"**Question {idx + 1}:** {item['question']}")
        st.sidebar.markdown(f"**Answer {idx + 1}:** {item['answer']}")
        unique_key = f"delete_favorite_{idx}_{time.time()}"
        delete_button = st.sidebar.button(f"Delete favorite item {idx + 1}", key=unique_key)
        if delete_button:
            st.session_state.favorites.pop(idx)
            display_history_and_favorites()
            break

def export_results(answer="", export_format="", file_name="", question=""):
    try:
        if export_format == "DOCX":
            doc = Document()
            doc.add_paragraph(question, style='Heading 1')
            doc.add_paragraph(answer, style='Normal')
            with io.BytesIO() as doc_bytes:
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                file_output = doc_bytes.read()

        b64 = base64.b64encode(file_output).decode()
        href = f'<a href="data:file/octet-stream;base64,{b64}" download="{file_name}.{export_format.lower()}">Download {file_name}.{export_format.lower()}</a>'
        st.markdown(href, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"An error occurred while exporting the results: {e}")

def main():
    openai.api_key = st.secrets["OPEN_API"]

    st.title("File Upload and GPT-4 Q&A")
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "txt", "png", "jpg", "jpeg"])

    if uploaded_file:
        document_text = process_uploaded_file(uploaded_file)
        if document_text:
            st.write("Document Content:")
            st.write(document_text)

            st.sidebar.header("GPT-4 Settings")
            temperature = st.sidebar.slider("Temperature", 0.1, 1.0, 0.5, 0.1)
            max_tokens = st.sidebar.slider("Max tokens", 10, 500, 150, 10)
            top_p = st.sidebar.slider("Top-p sampling", 0.0, 1.0, 1.0, 0.5)

            user_question = st.text_area("Ask a question about the document:")

            chunk_size = 4096
            chunks = [document_text[i:i+chunk_size] for i in range(0, len(document_text), chunk_size)]

            answer = ""

            if st.button("Get Answer"):
                for chunk in chunks:
                    prompt = f"Answer the following question based on the document's content:\n\n{chunk}\n\nQuestion: {user_question}\nAnswer:"
                    chunk_answer = generate_answer(prompt, temperature, max_tokens, top_p)
                    answer += chunk_answer
                st.write("Answer:")
                st.write(answer)

                st.session_state.history.append({"question": user_question, "answer": answer})
                export_results(answer=answer, export_format="DOCX", file_name="file1", question=user_question)

                if st.button("Add to favorites"):
                    st.session_state.favorites.append({"question": user_question, "answer": answer})
                    display_history_and_favorites()
        else:
            st.error("Unable to process the uploaded file")

    display_history_and_favorites()

if __name__ == "__main__":
    main()


